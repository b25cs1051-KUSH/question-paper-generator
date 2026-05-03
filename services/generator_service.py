import random
from database.db_manager import DatabaseManager
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class InsufficientQuestionsError(Exception):
    """Custom exception raised when a chapter lacks enough questions of a specific type."""
    def __init__(self, chapter_name, type_name, required, available):
        self.message = f"Chapter '{chapter_name}' has insufficient '{type_name}' questions. Required: {required}, Available: {available}."
        super().__init__(self.message)

class PaperGenerator:
    """
    Core engine for generating question papers based on templates and weights.
    Handles nested sections and blocks.
    """

    def __init__(self):
        self.db = DatabaseManager()

    def _get_chapter_name(self, chapter_id):
        conn = self.db._get_connection()
        try:
            row = conn.execute("SELECT name FROM chapters WHERE id = ?", (chapter_id,)).fetchone()
            return row['name'] if row else f"ID: {chapter_id}"
        finally:
            conn.close()

    def _get_type_name(self, type_id):
        conn = self.db._get_connection()
        try:
            row = conn.execute("SELECT name FROM question_types WHERE id = ?", (type_id,)).fetchone()
            return row['name'] if row else f"ID: {type_id}"
        finally:
            conn.close()

    def generate_paper(self, sections_config, seed):
        """
        Generates paper from a nested structure.
        sections_config: [{'name': 'Section A', 'blocks': [{'type_id': 1, 'chapters': [1,2], 'count': 5, 'marks': 1}]}]
        """
        random.seed(seed)
        paper = {
            "seed": seed,
            "sections": {}
        }

        for sec in sections_config:
            section_name = sec['name']
            paper["sections"][section_name] = {"blocks": []}
            
            for block in sec['blocks']:
                type_id = block['type_id']
                selected_chapters = block['chapters']
                total_required = block['count']
                
                if not selected_chapters:
                    continue

                # Weightage Logic
                num_chapters = len(selected_chapters)
                base_count = total_required // num_chapters
                remainder = total_required % num_chapters
                extra_chapters = random.sample(selected_chapters, remainder)
                
                block_questions = []
                for ch_id in selected_chapters:
                    count_for_this_chapter = base_count + (1 if ch_id in extra_chapters else 0)
                    if count_for_this_chapter == 0: continue

                    available_count = self.db.check_question_availability([ch_id], type_id)
                    if available_count < count_for_this_chapter:
                        raise InsufficientQuestionsError(self._get_chapter_name(ch_id), self._get_type_name(type_id), count_for_this_chapter, available_count)
                    
                    conn = self.db._get_connection()
                    try:
                        cursor = conn.execute("SELECT * FROM questions WHERE chapter_id = ? AND type_id = ?", (ch_id, type_id))
                        all_qs = [dict(row) for row in cursor.fetchall()]
                        block_questions.extend(random.sample(all_qs, count_for_this_chapter))
                    finally:
                        conn.close()

                random.shuffle(block_questions)
                paper["sections"][section_name]["blocks"].append({
                    "rule": block,
                    "questions": block_questions
                })

        return paper

class DocxExporter:
    @staticmethod
    def generate_docx(paper_data, output_path):
        doc = Document()
        
        # Header
        meta = paper_data.get('meta', {})
        title = doc.add_heading(meta.get('institution', 'ACADEMIC INSTITUTION').upper(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        exam = doc.add_heading(meta.get('exam_name', 'ANNUAL EXAMINATION').upper(), 1)
        exam.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Info Row
        p = doc.add_paragraph()
        p.add_run(f"Time: {meta.get('time_limit', '3 Hours')}").bold = True
        p.add_run("\t" * 4)
        total_marks = 0
        for sec in paper_data['sections'].values():
            for block in sec['blocks']:
                total_marks += block['rule']['count'] * block['rule']['marks']
        p.add_run(f"Total Marks: {total_marks}").bold = True
        
        doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

        q_num = 1
        for sec_name, sec_data in paper_data['sections'].items():
            sec_h = doc.add_heading(sec_name.upper(), 2)
            sec_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            for block in sec_data['blocks']:
                # Instruction
                instr = doc.add_paragraph()
                r = instr.add_run(f"Answer the following {block['rule'].get('count')} questions. ({block['rule'].get('marks')} Marks each)")
                r.italic = True
                r.bold = True
                
                for q in block['questions']:
                    q_p = doc.add_paragraph()
                    q_p.add_run(f"Q{q_num}. {q['content']}").font.size = Pt(11)
                    q_p.add_run(f"\t[{q['marks']}M]").bold = True
                    q_num += 1
                    
        doc.save(output_path)
